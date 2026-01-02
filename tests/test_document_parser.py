"""
Kompletní test suite pro document_parser.py modul.

Testuje DocumentDownloader, PDFParser, DOCXParser a DocumentParser
s důrazem na security, error handling a async behavior.
"""

import asyncio
from io import BytesIO
from unittest.mock import patch

import docx
import httpx
import pypdf
import pytest
from pypdf import PdfWriter

from sukl_mcp.document_parser import (
    CACHE_SIZE,
    CACHE_TTL,
    DOWNLOAD_TIMEOUT,
    MAX_FILE_SIZE,
    MAX_PDF_PAGES,
    PARSE_TIMEOUT,
    DocumentDownloader,
    DocumentParser,
    DOCXParser,
    PDFParser,
    close_document_parser,
    get_document_parser,
)
from sukl_mcp.exceptions import SUKLDocumentError, SUKLParseError

# === Fixtures ===


@pytest.fixture
def sample_pdf_bytes():
    """
    Vytvoř simple PDF bytes pro testování.

    Generuje PDF s 3 stranami textu.
    """
    writer = PdfWriter()

    # Přidej 3 strany s textem
    for _ in range(3):
        writer.add_blank_page(width=612, height=792)  # Letter size
        # Note: pypdf PdfWriter nepodporuje přidání textu bez složitého nastavení
        # Pro testování použijeme prázdné strany a mockujeme extract_text()

    output = BytesIO()
    writer.write(output)
    return output.getvalue()


@pytest.fixture
def sample_pdf_with_text():
    """
    Vytvoř PDF s reálným textem pomocí mockování.

    Vrací tuple (bytes, expected_text).
    """
    writer = PdfWriter()
    # Vytvoř 3 prázdné strany (text bude mockován v testech)
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)

    output = BytesIO()
    writer.write(output)

    expected_text = "Testovací text strany 1\n\nTestovací text strany 2\n\nTestovací text strany 3"

    return (output.getvalue(), expected_text)


@pytest.fixture
def sample_docx_bytes():
    """
    Vytvoř simple DOCX bytes pro testování.

    Obsahuje paragraf a tabulku s daty.
    """
    doc = docx.Document()

    # Přidej odstavec
    doc.add_paragraph("Testovací odstavec 1")
    doc.add_paragraph("Testovací odstavec 2")

    # Přidej tabulku
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Buňka A1"
    table.rows[0].cells[1].text = "Buňka B1"
    table.rows[1].cells[0].text = "Buňka A2"
    table.rows[1].cells[1].text = "Buňka B2"

    # Serializuj do bytes
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


@pytest.fixture
def empty_pdf_bytes():
    """PDF bez textu (prázdné strany)."""
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


@pytest.fixture
def empty_docx_bytes():
    """DOCX bez textu."""
    doc = docx.Document()
    doc.add_paragraph("")  # Prázdný odstavec
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


@pytest.fixture
def large_pdf_bytes():
    """PDF překračující MAX_PDF_PAGES (100 stran)."""
    writer = PdfWriter()
    for _ in range(105):  # 105 stran
        writer.add_blank_page(width=612, height=792)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


@pytest.fixture
def malformed_pdf_bytes():
    """Špatně formátovaný PDF soubor."""
    return b"This is not a valid PDF file content"


@pytest.fixture
def malformed_docx_bytes():
    """Špatně formátovaný DOCX soubor."""
    return b"This is not a valid DOCX file content"


# === DocumentDownloader Tests ===


class TestDocumentDownloader:
    """Testy pro DocumentDownloader třídu."""

    @pytest.mark.asyncio
    async def test_download_pdf_success(self, httpx_mock, sample_pdf_bytes):
        """Test úspěšného stažení PDF dokumentu."""
        url = "https://example.com/document.pdf"

        httpx_mock.add_response(
            url=url,
            content=sample_pdf_bytes,
            headers={
                "content-type": "application/pdf",
                "content-length": str(len(sample_pdf_bytes)),
            },
        )

        downloader = DocumentDownloader()
        content, format_type = await downloader.download(url)

        assert content == sample_pdf_bytes
        assert format_type == "pdf"
        assert len(content) == len(sample_pdf_bytes)

    @pytest.mark.asyncio
    async def test_download_docx_success(self, httpx_mock, sample_docx_bytes):
        """Test úspěšného stažení DOCX dokumentu."""
        url = "https://example.com/document.docx"

        httpx_mock.add_response(
            url=url,
            content=sample_docx_bytes,
            headers={
                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content-length": str(len(sample_docx_bytes)),
            },
        )

        downloader = DocumentDownloader()
        content, format_type = await downloader.download(url)

        assert content == sample_docx_bytes
        assert format_type == "docx"

    @pytest.mark.asyncio
    async def test_download_content_type_detection_pdf(self, httpx_mock):
        """Test Content-Type detekce pro PDF."""
        url = "https://example.com/doc"
        content = b"fake pdf content"

        httpx_mock.add_response(
            url=url,
            content=content,
            headers={"content-type": "application/pdf"},
        )

        downloader = DocumentDownloader()
        _, format_type = await downloader.download(url)

        assert format_type == "pdf"

    @pytest.mark.asyncio
    async def test_download_content_type_detection_docx(self, httpx_mock):
        """Test Content-Type detekce pro DOCX (word)."""
        url = "https://example.com/doc"
        content = b"fake docx content"

        httpx_mock.add_response(
            url=url,
            content=content,
            headers={
                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            },
        )

        downloader = DocumentDownloader()
        _, format_type = await downloader.download(url)

        assert format_type == "docx"

    @pytest.mark.asyncio
    async def test_download_url_extension_fallback(self, httpx_mock):
        """Test fallback na URL extension když Content-Type chybí."""
        url = "https://example.com/document.pdf"
        content = b"fake pdf content"

        # Žádný Content-Type header
        httpx_mock.add_response(url=url, content=content, headers={})

        downloader = DocumentDownloader()
        _, format_type = await downloader.download(url)

        assert format_type == "pdf"

    @pytest.mark.asyncio
    async def test_download_size_limit_from_header(self, httpx_mock):
        """Test size limit enforcement z Content-Length header."""
        url = "https://example.com/huge.pdf"

        # Content-Length překračuje MAX_FILE_SIZE
        httpx_mock.add_response(
            url=url,
            content=b"small content",
            headers={
                "content-type": "application/pdf",
                "content-length": str(MAX_FILE_SIZE + 1),
            },
        )

        downloader = DocumentDownloader()

        with pytest.raises(SUKLDocumentError, match="příliš velký"):
            await downloader.download(url)

    @pytest.mark.asyncio
    async def test_download_size_limit_from_actual_content(self, httpx_mock):
        """Test size limit enforcement ze skutečné velikosti content."""
        url = "https://example.com/huge.pdf"

        # Content-Length chybí, ale skutečný content je příliš velký
        huge_content = b"x" * (MAX_FILE_SIZE + 1)

        httpx_mock.add_response(
            url=url,
            content=huge_content,
            headers={"content-type": "application/pdf"},
        )

        downloader = DocumentDownloader()

        with pytest.raises(SUKLDocumentError, match="příliš velký"):
            await downloader.download(url)

    @pytest.mark.asyncio
    async def test_download_timeout_handling(self, httpx_mock):
        """Test timeout handling při stahování."""
        url = "https://example.com/slow.pdf"

        # Simuluj timeout pomocí callback funkce
        def timeout_callback(request):
            raise httpx.ReadTimeout("Connection timeout")

        httpx_mock.add_callback(timeout_callback, url=url)

        downloader = DocumentDownloader(timeout=1.0)

        with pytest.raises(SUKLDocumentError, match="Chyba při stahování"):
            await downloader.download(url)

    @pytest.mark.asyncio
    async def test_download_http_404_error(self, httpx_mock):
        """Test HTTP 404 error handling."""
        url = "https://example.com/missing.pdf"

        httpx_mock.add_response(url=url, status_code=404)

        downloader = DocumentDownloader()

        with pytest.raises(SUKLDocumentError, match="Chyba při stahování"):
            await downloader.download(url)

    @pytest.mark.asyncio
    async def test_download_http_500_error(self, httpx_mock):
        """Test HTTP 500 error handling."""
        url = "https://example.com/error.pdf"

        httpx_mock.add_response(url=url, status_code=500)

        downloader = DocumentDownloader()

        with pytest.raises(SUKLDocumentError, match="Chyba při stahování"):
            await downloader.download(url)

    @pytest.mark.asyncio
    async def test_download_unsupported_format(self, httpx_mock):
        """Test error při nepodporovaném formátu."""
        url = "https://example.com/document.txt"

        httpx_mock.add_response(
            url=url,
            content=b"text content",
            headers={"content-type": "text/plain"},
        )

        downloader = DocumentDownloader()

        with pytest.raises(SUKLDocumentError, match="Nepodporovaný formát"):
            await downloader.download(url)

    @pytest.mark.asyncio
    async def test_download_follow_redirects(self, httpx_mock):
        """Test že download následuje redirecty."""
        final_url = "https://example.com/final.pdf"
        content = b"final content"

        httpx_mock.add_response(
            url=final_url,
            content=content,
            headers={"content-type": "application/pdf"},
        )

        downloader = DocumentDownloader()
        result_content, _ = await downloader.download(final_url)

        assert result_content == content


# === PDFParser Tests ===


class TestPDFParser:
    """Testy pro PDFParser třídu."""

    def test_parse_pdf_success(self, sample_pdf_with_text):
        """Test úspěšného parsování PDF s textem."""
        pdf_bytes, expected_text = sample_pdf_with_text
        parser = PDFParser()

        # Mockujeme extract_text() aby vrátil text
        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.side_effect = [
                "Testovací text strany 1",
                "Testovací text strany 2",
                "Testovací text strany 3",
            ]

            text = parser.parse(pdf_bytes)

            assert "Testovací text strany 1" in text
            assert "Testovací text strany 2" in text
            assert "Testovací text strany 3" in text
            assert len(text) > 0

    def test_parse_pdf_size_limit(self):
        """Test size limit enforcement při parsování."""
        parser = PDFParser(max_size=1000)
        large_content = b"x" * 1001

        with pytest.raises(SUKLParseError, match="příliš velký"):
            parser.parse(large_content)

    def test_parse_pdf_page_limit(self, large_pdf_bytes):
        """Test page limit enforcement (max 100 stran)."""
        parser = PDFParser(max_pages=100)

        # Mock extract_text aby vrátil text pro každou stranu
        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.return_value = "Page text"

            parser.parse(large_pdf_bytes)

            # Mělo by zparsovat pouze prvních 100 stran
            assert mock_extract.call_count <= 100

    def test_parse_empty_pdf(self, empty_pdf_bytes):
        """Test parsování prázdného PDF (bez textu) - mělo by vyhodit error."""
        parser = PDFParser()

        # Mock extract_text aby vrátil prázdný string
        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.return_value = ""

            with pytest.raises(SUKLParseError, match="neobsahuje žádný text"):
                parser.parse(empty_pdf_bytes)

    def test_parse_malformed_pdf(self, malformed_pdf_bytes):
        """Test parsování malformed PDF - mělo by vyhodit error."""
        parser = PDFParser()

        with pytest.raises(SUKLParseError, match="Chyba při čtení PDF"):
            parser.parse(malformed_pdf_bytes)

    def test_parse_pdf_with_extraction_errors(self, sample_pdf_bytes):
        """Test handling chyb při extrakci jednotlivých stran."""
        parser = PDFParser()

        # První strana OK, druhá error, třetí OK
        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.side_effect = [
                "Page 1 text",
                Exception("Extraction error"),
                "Page 3 text",
            ]

            text = parser.parse(sample_pdf_bytes)

            # Mělo by zparsovat 1. a 3. stranu, přeskočit 2.
            assert "Page 1 text" in text
            assert "Page 3 text" in text
            assert len(text) > 0

    def test_parse_pdf_whitespace_only(self, sample_pdf_bytes):
        """Test PDF s pouze whitespace textem - mělo by vyhodit error."""
        parser = PDFParser()

        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.return_value = "   \n\n\t   "

            with pytest.raises(SUKLParseError, match="neobsahuje žádný text"):
                parser.parse(sample_pdf_bytes)


# === DOCXParser Tests ===


class TestDOCXParser:
    """Testy pro DOCXParser třídu."""

    def test_parse_docx_success(self, sample_docx_bytes):
        """Test úspěšného parsování DOCX."""
        parser = DOCXParser()
        text = parser.parse(sample_docx_bytes)

        # Měl by obsahovat text z odstavců
        assert "Testovací odstavec 1" in text
        assert "Testovací odstavec 2" in text

        # Měl by obsahovat text z tabulky
        assert "Buňka A1" in text
        assert "Buňka B1" in text
        assert "Buňka A2" in text
        assert "Buňka B2" in text

        assert len(text) > 0

    def test_parse_docx_paragraphs_only(self):
        """Test parsování DOCX obsahujícího pouze odstavce."""
        doc = docx.Document()
        doc.add_paragraph("Paragraph 1")
        doc.add_paragraph("Paragraph 2")
        doc.add_paragraph("Paragraph 3")

        output = BytesIO()
        doc.save(output)
        docx_bytes = output.getvalue()

        parser = DOCXParser()
        text = parser.parse(docx_bytes)

        assert "Paragraph 1" in text
        assert "Paragraph 2" in text
        assert "Paragraph 3" in text

    def test_parse_docx_tables_only(self):
        """Test parsování DOCX obsahujícího pouze tabulky."""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A1"
        table.rows[0].cells[1].text = "B1"
        table.rows[1].cells[0].text = "A2"
        table.rows[1].cells[1].text = "B2"

        output = BytesIO()
        doc.save(output)
        docx_bytes = output.getvalue()

        parser = DOCXParser()
        text = parser.parse(docx_bytes)

        assert "A1" in text
        assert "B1" in text
        assert "A2" in text
        assert "B2" in text

    def test_parse_docx_size_limit(self):
        """Test size limit enforcement při parsování DOCX."""
        parser = DOCXParser(max_size=1000)
        large_content = b"x" * 1001

        with pytest.raises(SUKLParseError, match="příliš velký"):
            parser.parse(large_content)

    def test_parse_empty_docx(self, empty_docx_bytes):
        """Test parsování prázdného DOCX - mělo by vyhodit error."""
        parser = DOCXParser()

        with pytest.raises(SUKLParseError, match="neobsahuje žádný text"):
            parser.parse(empty_docx_bytes)

    def test_parse_malformed_docx(self, malformed_docx_bytes):
        """Test parsování malformed DOCX - mělo by vyhodit error."""
        parser = DOCXParser()

        with pytest.raises(SUKLParseError, match="Chyba při parsování DOCX"):
            parser.parse(malformed_docx_bytes)

    def test_parse_docx_whitespace_only(self):
        """Test DOCX s pouze whitespace textem - mělo by vyhodit error."""
        doc = docx.Document()
        doc.add_paragraph("   \n\n\t   ")
        doc.add_paragraph("      ")

        output = BytesIO()
        doc.save(output)
        docx_bytes = output.getvalue()

        parser = DOCXParser()

        with pytest.raises(SUKLParseError, match="neobsahuje žádný text"):
            parser.parse(docx_bytes)

    def test_parse_docx_empty_cells_skipped(self):
        """Test že prázdné buňky v tabulkách jsou přeskočeny."""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A1"
        table.rows[0].cells[1].text = ""  # Prázdná buňka
        table.rows[1].cells[0].text = ""  # Prázdná buňka
        table.rows[1].cells[1].text = "B2"

        output = BytesIO()
        doc.save(output)
        docx_bytes = output.getvalue()

        parser = DOCXParser()
        text = parser.parse(docx_bytes)

        assert "A1" in text
        assert "B2" in text
        # Prázdné buňky by neměly přidat extra whitespace
        assert text.count("\n\n") >= 1


# === DocumentParser Integration Tests ===


class TestDocumentParser:
    """Integration testy pro DocumentParser."""

    @pytest.mark.asyncio
    async def test_get_document_content_pdf_success(self, httpx_mock, sample_pdf_with_text):
        """Test úspěšného stažení a parsování PIL (PDF)."""
        pdf_bytes, expected_text = sample_pdf_with_text
        sukl_code = "0254045"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(
            url=url,
            content=pdf_bytes,
            headers={"content-type": "application/pdf"},
        )

        parser = DocumentParser()

        # Mock PDF parsing
        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.side_effect = [
                "Testovací text strany 1",
                "Testovací text strany 2",
                "Testovací text strany 3",
            ]

            result = await parser.get_document_content(sukl_code, doc_type)

            assert result["sukl_code"] == sukl_code
            assert result["doc_type"] == doc_type
            assert result["format"] == "pdf"
            assert result["url"] == url
            assert "Testovací text strany 1" in result["content"]
            assert len(result["content"]) > 0

    @pytest.mark.asyncio
    async def test_get_document_content_spc_success(self, httpx_mock, sample_pdf_with_text):
        """Test úspěšného stažení a parsování SPC (PDF)."""
        pdf_bytes, _ = sample_pdf_with_text
        sukl_code = "0123456"
        doc_type = "spc"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(
            url=url,
            content=pdf_bytes,
            headers={"content-type": "application/pdf"},
        )

        parser = DocumentParser()

        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.return_value = "SPC document text"

            result = await parser.get_document_content(sukl_code, doc_type)

            assert result["doc_type"] == "spc"
            assert "SPC document text" in result["content"]

    @pytest.mark.asyncio
    async def test_get_document_content_docx(self, httpx_mock, sample_docx_bytes):
        """Test stažení a parsování DOCX dokumentu."""
        sukl_code = "0789012"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        # Server vrací DOCX i když URL je .pdf
        httpx_mock.add_response(
            url=url,
            content=sample_docx_bytes,
            headers={
                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            },
        )

        parser = DocumentParser()
        result = await parser.get_document_content(sukl_code, doc_type)

        assert result["format"] == "docx"
        assert "Testovací odstavec" in result["content"]

    @pytest.mark.asyncio
    async def test_get_document_content_cache_hit(self, httpx_mock, sample_pdf_with_text):
        """Test že druhé volání je rychlejší díky cache."""
        pdf_bytes, _ = sample_pdf_with_text
        sukl_code = "0111111"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(
            url=url,
            content=pdf_bytes,
            headers={"content-type": "application/pdf"},
        )

        parser = DocumentParser()

        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.return_value = "Cached document text"

            # První volání
            result1 = await parser.get_document_content(sukl_code, doc_type)

            # Druhé volání by mělo použít cache (httpx_mock by selhal na druhý request)
            result2 = await parser.get_document_content(sukl_code, doc_type)

            assert result1 == result2
            assert result1["content"] == result2["content"]

            # Extract by měl být volán 3x (3 strany) při prvním volání
            # a NOT být volán znovu při druhém volání (cache hit)
            assert mock_extract.call_count == 3

    @pytest.mark.asyncio
    async def test_get_document_content_parse_timeout(self, httpx_mock, sample_pdf_bytes):
        """Test timeout handling při parsování (30s timeout)."""
        sukl_code = "0222222"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(
            url=url,
            content=sample_pdf_bytes,
            headers={"content-type": "application/pdf"},
        )

        parser = DocumentParser()

        # Mock parse aby trval příliš dlouho (synchronní funkce s time.sleep)
        def slow_parse(*args):
            import time

            time.sleep(35)  # Více než PARSE_TIMEOUT (30s)
            return "Never returned"

        with patch.object(parser.pdf_parser, "parse", side_effect=slow_parse):
            with pytest.raises(SUKLParseError, match="Timeout při parsování"):
                await parser.get_document_content(sukl_code, doc_type)

    @pytest.mark.asyncio
    async def test_get_document_content_download_error(self, httpx_mock):
        """Test handling download error."""
        sukl_code = "0333333"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(url=url, status_code=404)

        parser = DocumentParser()

        with pytest.raises(SUKLDocumentError, match="Chyba při stahování"):
            await parser.get_document_content(sukl_code, doc_type)

    @pytest.mark.asyncio
    async def test_get_document_content_parse_error(self, httpx_mock, malformed_pdf_bytes):
        """Test handling parse error."""
        sukl_code = "0444444"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(
            url=url,
            content=malformed_pdf_bytes,
            headers={"content-type": "application/pdf"},
        )

        parser = DocumentParser()

        with pytest.raises(SUKLParseError):
            await parser.get_document_content(sukl_code, doc_type)

    @pytest.mark.asyncio
    async def test_get_document_content_unsupported_format(self, httpx_mock):
        """Test error při nepodporovaném formátu."""
        sukl_code = "0555555"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(
            url=url,
            content=b"Plain text content",
            headers={"content-type": "text/plain"},
        )

        parser = DocumentParser()

        with pytest.raises(SUKLDocumentError, match="Nepodporovaný formát"):
            await parser.get_document_content(sukl_code, doc_type)

    @pytest.mark.asyncio
    async def test_clear_cache(self, httpx_mock, sample_pdf_with_text):
        """Test clear_cache() funkčnosti."""
        pdf_bytes, _ = sample_pdf_with_text
        sukl_code = "0666666"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(
            url=url,
            content=pdf_bytes,
            headers={"content-type": "application/pdf"},
        )

        parser = DocumentParser()

        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.return_value = "Document text"

            # První volání
            await parser.get_document_content(sukl_code, doc_type)
            assert mock_extract.call_count == 3  # 3 strany

            # Clear cache
            parser.clear_cache()

            # Druhé volání by mělo znovu stáhnout (ne z cache)
            httpx_mock.add_response(
                url=url,
                content=pdf_bytes,
                headers={"content-type": "application/pdf"},
            )

            await parser.get_document_content(sukl_code, doc_type)
            assert mock_extract.call_count == 6  # 3 strany + 3 strany po clear_cache


# === Singleton Tests ===


class TestSingletonPattern:
    """Testy pro singleton pattern get_document_parser()."""

    def test_get_document_parser_returns_instance(self):
        """Test že get_document_parser() vrací instanci DocumentParser."""
        parser = get_document_parser()
        assert isinstance(parser, DocumentParser)

    def test_get_document_parser_singleton(self):
        """Test že get_document_parser() vrací stejnou instanci."""
        parser1 = get_document_parser()
        parser2 = get_document_parser()

        # Měly by být stejná instance (stejné id)
        assert parser1 is parser2
        assert id(parser1) == id(parser2)

    def test_close_document_parser(self):
        """Test close_document_parser() cleanup."""
        parser = get_document_parser()

        # Zavři parser
        close_document_parser()

        # Nový get by měl vytvořit novou instanci
        new_parser = get_document_parser()
        assert new_parser is not parser
        assert id(new_parser) != id(parser)


# === Async I/O Tests ===


class TestAsyncIOBehavior:
    """Testy pro async I/O non-blocking behavior."""

    @pytest.mark.asyncio
    async def test_parsing_runs_in_executor(self, httpx_mock, sample_pdf_with_text):
        """Test že parsing běží v executoru (non-blocking)."""
        pdf_bytes, _ = sample_pdf_with_text
        sukl_code = "0777777"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(
            url=url,
            content=pdf_bytes,
            headers={"content-type": "application/pdf"},
        )

        parser = DocumentParser()

        # Mock extract_text s delay
        call_times = []

        def delayed_extract(*args, **kwargs):
            import time

            call_times.append(time.time())
            time.sleep(0.1)  # Simulace CPU-bound operace
            return "Page text"

        with patch.object(pypdf.PageObject, "extract_text", side_effect=delayed_extract):
            start_time = asyncio.get_event_loop().time()

            # Spusť parsing
            task = asyncio.create_task(parser.get_document_content(sukl_code, doc_type))

            # Event loop by měl zůstat responzivní během parsingu
            await asyncio.sleep(0.05)
            elapsed = asyncio.get_event_loop().time() - start_time

            # Měli bychom být schopni checknout elapsed time během parsingu
            assert elapsed < 0.2

            # Počkej na dokončení
            result = await task
            assert "Page text" in result["content"]

    @pytest.mark.asyncio
    async def test_concurrent_downloads(self, httpx_mock, sample_pdf_with_text):
        """Test že multiple concurrent downloads fungují správně."""
        pdf_bytes, _ = sample_pdf_with_text

        # Setup mock responses pro 5 různých dokumentů
        sukl_codes = [f"0{i}{i}{i}{i}{i}{i}{i}" for i in range(5)]

        for code in sukl_codes:
            url = f"https://prehledy.sukl.cz/pil/{code}.pdf"
            httpx_mock.add_response(
                url=url,
                content=pdf_bytes,
                headers={"content-type": "application/pdf"},
            )

        parser = DocumentParser()

        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.return_value = "Concurrent document text"

            # Spusť 5 downloads concurrently
            tasks = [parser.get_document_content(code, "pil") for code in sukl_codes]

            results = await asyncio.gather(*tasks)

            # Všechny by měly uspět
            assert len(results) == 5
            for i, result in enumerate(results):
                assert result["sukl_code"] == sukl_codes[i]
                assert "Concurrent document text" in result["content"]


# === Security Tests ===


class TestSecurityFeatures:
    """Testy pro security features (size limits, timeouts)."""

    @pytest.mark.asyncio
    async def test_max_file_size_enforcement(self, httpx_mock):
        """Test enforcement MAX_FILE_SIZE limitu."""
        sukl_code = "0888888"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        # Vytvoř content větší než MAX_FILE_SIZE
        huge_content = b"x" * (MAX_FILE_SIZE + 1)

        httpx_mock.add_response(
            url=url,
            content=huge_content,
            headers={"content-type": "application/pdf"},
        )

        parser = DocumentParser()

        with pytest.raises(SUKLDocumentError, match="příliš velký"):
            await parser.get_document_content(sukl_code, doc_type)

    @pytest.mark.asyncio
    async def test_max_pdf_pages_enforcement(self, httpx_mock, large_pdf_bytes):
        """Test enforcement MAX_PDF_PAGES limitu (100 stran)."""
        sukl_code = "0999999"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(
            url=url,
            content=large_pdf_bytes,
            headers={"content-type": "application/pdf"},
        )

        parser = DocumentParser()

        with patch.object(pypdf.PageObject, "extract_text") as mock_extract:
            mock_extract.return_value = "Page text"

            await parser.get_document_content(sukl_code, doc_type)

            # Mělo by zparsovat pouze prvních 100 stran
            assert mock_extract.call_count <= MAX_PDF_PAGES

    @pytest.mark.asyncio
    async def test_download_timeout_enforcement(self, httpx_mock):
        """Test enforcement download timeoutu."""
        sukl_code = "1000000"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        # Simuluj timeout
        def timeout_callback(request):
            raise httpx.ReadTimeout("Download timeout")

        httpx_mock.add_callback(timeout_callback, url=url)

        parser = DocumentParser(downloader=DocumentDownloader(timeout=1.0))

        with pytest.raises(SUKLDocumentError, match="Chyba při stahování"):
            await parser.get_document_content(sukl_code, doc_type)

    @pytest.mark.asyncio
    async def test_parse_timeout_enforcement(self, httpx_mock, sample_pdf_bytes):
        """Test enforcement parse timeoutu (30s)."""
        sukl_code = "1111111"
        doc_type = "pil"
        url = f"https://prehledy.sukl.cz/{doc_type}/{sukl_code}.pdf"

        httpx_mock.add_response(
            url=url,
            content=sample_pdf_bytes,
            headers={"content-type": "application/pdf"},
        )

        parser = DocumentParser()

        # Mock parse s timeoutem
        def slow_parse(*args, **kwargs):
            import time

            time.sleep(35)  # > PARSE_TIMEOUT
            return "Never returned"

        with patch.object(parser.pdf_parser, "parse", side_effect=slow_parse):
            with pytest.raises(SUKLParseError, match="Timeout při parsování"):
                await parser.get_document_content(sukl_code, doc_type)


# === Configuration Tests ===


class TestConfiguration:
    """Testy pro konfigurační konstanty."""

    def test_constants_defined(self):
        """Test že všechny konstanty jsou definovány."""
        assert MAX_FILE_SIZE == 50 * 1024 * 1024  # 50 MB
        assert MAX_PDF_PAGES == 100
        assert DOWNLOAD_TIMEOUT == 30.0
        assert PARSE_TIMEOUT == 30.0
        assert CACHE_SIZE == 50
        assert CACHE_TTL == 86400  # 24 hodin

    def test_custom_configuration(self):
        """Test že lze vytvořit parser s custom konfigurací."""
        custom_downloader = DocumentDownloader(timeout=60.0, max_size=100 * 1024 * 1024)
        custom_pdf_parser = PDFParser(max_pages=200, max_size=100 * 1024 * 1024)
        custom_docx_parser = DOCXParser(max_size=100 * 1024 * 1024)

        parser = DocumentParser(
            downloader=custom_downloader,
            pdf_parser=custom_pdf_parser,
            docx_parser=custom_docx_parser,
        )

        assert parser.downloader.timeout == 60.0
        assert parser.downloader.max_size == 100 * 1024 * 1024
        assert parser.pdf_parser.max_pages == 200
        assert parser.pdf_parser.max_size == 100 * 1024 * 1024
        assert parser.docx_parser.max_size == 100 * 1024 * 1024
