"""
Document parser pro SÚKL dokumenty (PIL/SPC).

Podporuje stahování a parsing PDF a DOCX souborů s async cachingem.
"""

import asyncio
import logging
from io import BytesIO
from typing import Any, Literal

import docx
import httpx
import pypdf
from async_lru import alru_cache

# Absolutní importy pro FastMCP Cloud compatibility
from sukl_mcp.exceptions import SUKLDocumentError, SUKLParseError

logger = logging.getLogger(__name__)

# === Configuration ===

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
MAX_PDF_PAGES = 100  # Limit pro prevenci DoS
DOWNLOAD_TIMEOUT = 30.0  # 30s timeout pro download
PARSE_TIMEOUT = 30.0  # 30s timeout pro parsing
CACHE_SIZE = 50  # Max 50 dokumentů v cache
CACHE_TTL = 86400  # 24 hodin TTL


# === Document Downloader ===


class DocumentDownloader:
    """Async downloader pro PIL/SPC dokumenty."""

    def __init__(
        self,
        timeout: float = DOWNLOAD_TIMEOUT,
        max_size: int = MAX_FILE_SIZE,
    ):
        self.timeout = timeout
        self.max_size = max_size

    async def download(self, url: str) -> tuple[bytes, str]:
        """
        Stáhni dokument z URL.

        Args:
            url: URL dokumentu (PDF nebo DOCX)

        Returns:
            Tuple (content_bytes, format_type)
            format_type je "pdf" nebo "docx"

        Raises:
            SUKLDocumentError: Když download selže nebo soubor je příliš velký
        """
        logger.info(f"Stahuji dokument: {url}")

        try:
            async with httpx.AsyncClient(timeout=self.timeout) as client:
                response = await client.get(url, follow_redirects=True)
                response.raise_for_status()

                # Kontrola velikosti
                content_length = response.headers.get("content-length")
                if content_length and int(content_length) > self.max_size:
                    raise SUKLDocumentError(
                        f"Dokument příliš velký: {int(content_length)} B (max {self.max_size} B)"
                    )

                content = response.content

                # Kontrola velikosti skutečného obsahu
                if len(content) > self.max_size:
                    raise SUKLDocumentError(
                        f"Dokument příliš velký: {len(content)} B (max {self.max_size} B)"
                    )

                # Detekce formátu z Content-Type (má přednost před URL)
                content_type = response.headers.get("content-type", "").lower()

                # Pokud máme validní Content-Type (ne prázdný a ne generický), použij ho
                if content_type and content_type != "application/octet-stream":
                    if "pdf" in content_type:
                        format_type = "pdf"
                    elif "word" in content_type or "docx" in content_type:
                        format_type = "docx"
                    else:
                        # Nepodporovaný Content-Type
                        raise SUKLDocumentError(
                            f"Nepodporovaný formát dokumentu: {content_type}"
                        )
                # Fallback na URL extension pokud Content-Type chybí nebo je generický
                elif url.lower().endswith(".pdf"):
                    format_type = "pdf"
                elif url.lower().endswith(".docx"):
                    format_type = "docx"
                else:
                    raise SUKLDocumentError(
                        "Nepodporovaný formát dokumentu (neznámá URL extension)"
                    )

                logger.info(f"Dokument stažen: {len(content)} B, formát: {format_type}")
                return (content, format_type)

        except httpx.HTTPError as e:
            raise SUKLDocumentError(f"Chyba při stahování dokumentu: {e}") from e
        except Exception as e:
            raise SUKLDocumentError(f"Neočekávaná chyba při stahování: {e}") from e


# === PDF Parser ===


class PDFParser:
    """Parser pro PDF dokumenty s bezpečnostními limity."""

    def __init__(
        self,
        max_pages: int = MAX_PDF_PAGES,
        max_size: int = MAX_FILE_SIZE,
    ):
        self.max_pages = max_pages
        self.max_size = max_size

    def parse(self, content: bytes) -> str:
        """
        Parsuj PDF obsah do textu (synchronní).

        Args:
            content: PDF soubor jako bytes

        Returns:
            Extrahovaný text

        Raises:
            SUKLParseError: Když parsing selže
        """
        if len(content) > self.max_size:
            raise SUKLParseError(
                f"PDF příliš velký: {len(content)} B (max {self.max_size} B)"
            )

        try:
            reader = pypdf.PdfReader(BytesIO(content))

            # Kontrola počtu stran
            num_pages = len(reader.pages)
            if num_pages > self.max_pages:
                logger.warning(
                    f"PDF má {num_pages} stran, parsuju pouze prvních {self.max_pages}"
                )

            # Extrakce textu
            text_parts = []
            for page_num, page in enumerate(reader.pages[: self.max_pages]):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Chyba při extrakci strany {page_num}: {e}")
                    continue

            text = "\n\n".join(text_parts)

            if not text.strip():
                raise SUKLParseError("PDF neobsahuje žádný text")

            logger.info(f"PDF úspěšně parsováno: {len(text)} znaků z {num_pages} stran")
            return text

        except pypdf.errors.PdfReadError as e:
            raise SUKLParseError(f"Chyba při čtení PDF: {e}") from e
        except Exception as e:
            raise SUKLParseError(f"Neočekávaná chyba při parsování PDF: {e}") from e


# === DOCX Parser ===


class DOCXParser:
    """Parser pro DOCX dokumenty s bezpečnostními limity."""

    def __init__(self, max_size: int = MAX_FILE_SIZE):
        self.max_size = max_size

    def parse(self, content: bytes) -> str:
        """
        Parsuj DOCX obsah do textu (synchronní).

        Args:
            content: DOCX soubor jako bytes

        Returns:
            Extrahovaný text

        Raises:
            SUKLParseError: Když parsing selže
        """
        if len(content) > self.max_size:
            raise SUKLParseError(
                f"DOCX příliš velký: {len(content)} B (max {self.max_size} B)"
            )

        try:
            doc = docx.Document(BytesIO(content))

            # Extrakce textu z paragrafů
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extrakce textu z tabulek
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = "\n\n".join(text_parts)

            if not text.strip():
                raise SUKLParseError("DOCX neobsahuje žádný text")

            logger.info(f"DOCX úspěšně parsováno: {len(text)} znaků")
            return text

        except Exception as e:
            raise SUKLParseError(f"Chyba při parsování DOCX: {e}") from e


# === Main Document Parser ===


class DocumentParser:
    """
    Hlavní parser pro SÚKL dokumenty s async cachingem.

    Podporuje PDF a DOCX formáty s automatickou detekcí.
    """

    def __init__(
        self,
        downloader: DocumentDownloader | None = None,
        pdf_parser: PDFParser | None = None,
        docx_parser: DOCXParser | None = None,
    ):
        self.downloader = downloader or DocumentDownloader()
        self.pdf_parser = pdf_parser or PDFParser()
        self.docx_parser = docx_parser or DOCXParser()

    @alru_cache(maxsize=CACHE_SIZE, ttl=CACHE_TTL)
    async def get_document_content(
        self,
        sukl_code: str,
        doc_type: Literal["pil", "spc"],
    ) -> dict[str, Any]:
        """
        Stáhni a parsuj dokument (PIL nebo SPC) s cachingem.

        Cache: 50 dokumentů, 24h TTL

        Args:
            sukl_code: SÚKL kód léčiva (např. "0254045")
            doc_type: Typ dokumentu ("pil" nebo "spc")

        Returns:
            Dict s klíči:
            - content: Extrahovaný text
            - format: "pdf" nebo "docx"
            - url: URL dokumentu
            - sukl_code: SÚKL kód
            - doc_type: Typ dokumentu

        Raises:
            SUKLDocumentError: Když download selže
            SUKLParseError: Když parsing selže
        """
        # Konstruuj URL podle SÚKL formátu
        # Příklad: https://prehledy.sukl.cz/pil/0254045.pdf
        base_url = "https://prehledy.sukl.cz"
        url = f"{base_url}/{doc_type.lower()}/{sukl_code}.pdf"

        logger.info(f"Získávám dokument: {doc_type.upper()} pro {sukl_code}")

        try:
            # Download dokumentu
            content, format_type = await self.downloader.download(url)

            # Parsování v executoru (non-blocking)
            loop = asyncio.get_event_loop()

            if format_type == "pdf":
                # PDF parsing s timeoutem
                text = await asyncio.wait_for(
                    loop.run_in_executor(None, self.pdf_parser.parse, content),
                    timeout=PARSE_TIMEOUT,
                )
            elif format_type == "docx":
                # DOCX parsing s timeoutem
                text = await asyncio.wait_for(
                    loop.run_in_executor(None, self.docx_parser.parse, content),
                    timeout=PARSE_TIMEOUT,
                )
            else:
                raise SUKLDocumentError(f"Nepodporovaný formát: {format_type}")

            logger.info(
                f"Dokument úspěšně zpracován: {len(text)} znaků ({format_type})"
            )

            return {
                "content": text,
                "format": format_type,
                "url": url,
                "sukl_code": sukl_code,
                "doc_type": doc_type,
            }

        except asyncio.TimeoutError:
            raise SUKLParseError(
                f"Timeout při parsování dokumentu ({PARSE_TIMEOUT}s)"
            ) from None
        except (SUKLDocumentError, SUKLParseError):
            raise
        except Exception as e:
            raise SUKLDocumentError(
                f"Neočekávaná chyba při zpracování dokumentu: {e}"
            ) from e

    def clear_cache(self) -> None:
        """Vyčisti cache."""
        self.get_document_content.cache_clear()
        logger.info("Document cache vyčištěna")


# === Singleton instance ===

_parser: DocumentParser | None = None


def get_document_parser() -> DocumentParser:
    """
    Získej singleton instanci DocumentParser.

    Returns:
        DocumentParser instance
    """
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser


def close_document_parser() -> None:
    """Zavři a vyčisti parser (pro cleanup)."""
    global _parser
    if _parser is not None:
        _parser.clear_cache()
        _parser = None
        logger.info("DocumentParser zavřen")
